#!/usr/bin/env python


"""Match."""

from uuid import uuid4
import md5
from docx import Document
from docx.shared import Pt
import retinasdk
from retinasdk.client.exceptions import CorticalioException
from os import environ, linesep
from sqlalchemy import *

FILES_EXTENTION = ".docx"
RESULTS_FOLDER = "static/results_files"
DOCUMENT_HEADING = "Semantic matching results"
RETINA_API_KEY = environ['SME_RETINA_API_KEY']
RET_API_SERVER = "http://api.cortical.io/rest"
RET_NAME = "en_associative"
SME_SQLALCHEMY_DATABASE_URI = environ['SME_DEV_SQLALCHEMY_DATABASE_URI']


def get_clean_text(raw_text):
    """Clean text."""
    return "".join([el for el in raw_text.split(linesep)])

UNI_TO_ASCII_MAP = {u'\u2019': "'", u'\u2018': "`", u'\u2022': "_"}


def uni_to_ascii(s_in):
    """Unicode to ascii."""
    try:
        return str(s_in)
    except:
        pass
    s_out = ""
    for i in s_in:
        try:
            s_out = s_out + str(i)
        except:
            if i in UNI_TO_ASCII_MAP:
                s_out = s_out + UNI_TO_ASCII_MAP[i]
            else:
                s_out = s_out + "_"
    return s_out


def match_texts(text_1, text_2):
    """Match text."""
    text_1_raw = text_1
    text_2_raw = text_2
    text_1 = uni_to_ascii(text_1)
    text_2 = uni_to_ascii(text_2)
    uuid_str = str(uuid4())
    liteClient = retinasdk.LiteClient(RETINA_API_KEY)
    text_1_md5 = md5.new(text_1).hexdigest()
    text_2_md5 = md5.new(text_2).hexdigest()
    sim, text_1_fp, text_2_fp = "", "", ""
    text_1_keywords, text_2_keywords = "", ""
    try:
        text_1_fp = liteClient.getFingerprint(text_1)
    except CorticalioException:
        print(CorticalioException + "\nText:" + text_1)
        pass
    try:
        text_2_fp = liteClient.getFingerprint(text_2)
    except CorticalioException:
        print(str(CorticalioException) + "\nText:" + text_2)
        pass
    try:
        sim = str(liteClient.compare(text_1_fp, text_2_fp))
    except CorticalioException:
        print(CorticalioException)
        pass
    try:
        text_1_keywords = liteClient.getKeywords(text_1)
    except CorticalioException:
        print(CorticalioException)
        pass
    try:
        text_2_keywords = liteClient.getKeywords(text_2)
    except CorticalioException:
        print(CorticalioException)
        pass

    if len(text_1_keywords) == 0:
        text_1_keywords = ["[None]"]
    if len(text_2_keywords) == 0:
        text_2_keywords = ["[None]"]
    if text_1_fp == "":
        text_1_fp = ["[None]"]
    if text_2_fp == "":
        text_2_fp = ["[None]"]
    if sim == "" or text_1_fp == "[None]" or text_2_fp == "[None]":
        sim = "[None]"
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    document.add_heading(DOCUMENT_HEADING, 0)
    document.add_paragraph(
        'Text 1 and Text 2 similarity', style='Intense Quote')
    document.add_paragraph(sim)
    document.add_paragraph('Matching details', style='Intense Quote')
    table = document.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    font.bold = True
    table.cell(0, 1).text = "Text 1"
    table.cell(0, 2).text = "Text 2"
    table.cell(1, 0).text = "md5"
    table.cell(2, 0).text = "Keywords"
    table.cell(3, 0).text = "Fingerprint"
    font.bold = False
    table.cell(1, 1).text = text_1_md5
    table.cell(1, 2).text = text_2_md5
    table.cell(2, 1).text = ", ".join(text_1_keywords)
    table.cell(2, 2).text = ", ".join(text_2_keywords)
    table.cell(3, 1).text = ", ".join([str(el) for el in text_1_fp])
    table.cell(3, 2).text = ", ".join([str(el) for el in text_2_fp])
    document.add_page_break()
    document.add_heading('Original Text 1', 0)
    document.add_paragraph(get_clean_text(text_1_raw))
    document.add_page_break()
    document.add_heading('Original Text 2', 0)
    document.add_paragraph(get_clean_text(text_2_raw))
    document.save(RESULTS_FOLDER + "/" + uuid_str + FILES_EXTENTION)

    db = create_engine(SME_SQLALCHEMY_DATABASE_URI)
    db.echo = False
    metadata = MetaData(db)
    text_table = Table('text', metadata, autoload=True)
    text_table.insert().execute(text_md5=text_1_md5,
                                text_keywords=", ".join(text_1_keywords),
                                text_fingerprint=", ".join(
                                    [str(el) for el in text_1_fp]),
                                text=text_1)
    text_table.insert().execute(text_md5=text_2_md5,
                                text_keywords=", ".join(text_2_keywords),
                                text_fingerprint=", ".join(
                                    [str(el) for el in text_2_fp]),
                                text=text_2)

    return uuid_str
